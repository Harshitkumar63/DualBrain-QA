"""
rag_pipeline.py — RAG (Retrieval-Augmented Generation) Pipeline
================================================================

Upgrade features:
  1. Multi-format ingestion (PDF, DOCX, TXT, MD).
  2. Detailed metadata extraction.
  3. Hybrid Retrieval (ChromaDB + BM25 Sparse Search) with Reciprocal Rank Fusion (RRF).
  4. Query rewriting.
  5. Source citation extraction.
"""

from __future__ import annotations

import io
import logging
import os
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

import numpy as np
from pypdf import PdfReader
from docx import Document as DocxDocument
from rank_bm25 import BM25Okapi

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from src.chroma_store import ChromaStore

logger = logging.getLogger(__name__)

# ------------------------------------------------------------------ #
#  Configuration defaults                                             #
# ------------------------------------------------------------------ #

DEFAULT_EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
DEFAULT_CHUNK_SIZE = 512
DEFAULT_CHUNK_OVERLAP = 64
DEFAULT_TOP_K = 4
DEFAULT_PERSIST_DIR = "./data/chromadb"
DEFAULT_COLLECTION = "hybrid_rag"


class RAGPipeline:
    """End-to-end Hybrid RAG Pipeline using ChromaDB and BM25."""

    def __init__(
        self,
        embedding_model_name: str = DEFAULT_EMBEDDING_MODEL,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        top_k: int = DEFAULT_TOP_K,
        persist_dir: str = DEFAULT_PERSIST_DIR,
        collection_name: str = DEFAULT_COLLECTION,
    ) -> None:
        self.top_k = top_k
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # 1. Load Embedding Model
        logger.info("Loading embedding model: %s", embedding_model_name)
        self.embeddings = HuggingFaceEmbeddings(
            model_name=embedding_model_name,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )

        # 2. Text Splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

        # 3. Vector Database (ChromaDB Migration)
        self.vector_store = ChromaStore(
            persist_directory=persist_dir,
            collection_name=collection_name,
            embedding_model=self.embeddings,
        )

        # 4. Sparse Retriever State (BM25)
        self.bm25: Optional[BM25Okapi] = None
        self.bm25_docs: List[Document] = []
        
        # Build index at startup if database contains documents
        self.rebuild_bm25()

    # -------------------------------------------------------------- #
    #  Sparse Retriever Rebuild                                        #
    # -------------------------------------------------------------- #

    def rebuild_bm25(self) -> None:
        """Fetch all documents from ChromaDB and construct the BM25 index."""
        logger.info("Rebuilding BM25 keyword index...")
        try:
            db_data = self.vector_store._vector_store.get()
            documents: List[Document] = []
            if db_data and "documents" in db_data and db_data["documents"]:
                for i in range(len(db_data["documents"])):
                    content = db_data["documents"][i]
                    metadata = db_data["metadatas"][i] if db_data["metadatas"] else {}
                    documents.append(Document(page_content=content, metadata=metadata))

            if not documents:
                self.bm25 = None
                self.bm25_docs = []
                logger.info("BM25 index is empty (no documents in ChromaDB).")
                return

            self.bm25_docs = documents
            # Basic tokenization
            tokenized_corpus = [doc.page_content.lower().split() for doc in documents]
            self.bm25 = BM25Okapi(tokenized_corpus)
            logger.info("BM25 index successfully rebuilt with %d chunks.", len(documents))
        except Exception as e:
            logger.error("Failed to rebuild BM25 index: %s", e, exc_info=True)
            self.bm25 = None
            self.bm25_docs = []

    # -------------------------------------------------------------- #
    #  Document Parsing                                                #
    # -------------------------------------------------------------- #

    def parse_file(self, file_path: Path) -> List[Document]:
        """Parse text and metadata from a local file on disk."""
        filename = file_path.name
        ext = file_path.suffix.lower()
        timestamp = datetime.now(timezone.utc).isoformat()
        docs: List[Document] = []

        if ext == ".pdf":
            try:
                reader = PdfReader(file_path)
                for idx, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        docs.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "filename": filename,
                                    "page_number": idx + 1,
                                    "document_type": "pdf",
                                    "ingestion_timestamp": timestamp,
                                    "source": str(file_path),
                                },
                            )
                        )
            except Exception as e:
                logger.error("Failed to parse PDF %s: %s", filename, e)
        elif ext == ".docx":
            try:
                doc = DocxDocument(file_path)
                full_text = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n".join(full_text)
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={
                                "filename": filename,
                                "page_number": 1,
                                "document_type": "docx",
                                "ingestion_timestamp": timestamp,
                                "source": str(file_path),
                            },
                        )
                    )
            except Exception as e:
                logger.error("Failed to parse DOCX %s: %s", filename, e)
        elif ext in (".md", ".markdown", ".txt"):
            try:
                text = file_path.read_text(encoding="utf-8", errors="replace")
                doc_type = "markdown" if ext in (".md", ".markdown") else "txt"
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={
                                "filename": filename,
                                "page_number": 1,
                                "document_type": doc_type,
                                "ingestion_timestamp": timestamp,
                                "source": str(file_path),
                            },
                        )
                    )
            except Exception as e:
                logger.error("Failed to parse text file %s: %s", filename, e)
        else:
            logger.warning("Unsupported file type skipped: %s", filename)

        return docs

    def parse_file_bytes(self, file_bytes: bytes, filename: str) -> List[Document]:
        """Parse text and metadata from uploaded file bytes (for FastAPI upload)."""
        ext = os.path.splitext(filename)[1].lower()
        timestamp = datetime.now(timezone.utc).isoformat()
        docs: List[Document] = []

        if ext == ".pdf":
            try:
                f = io.BytesIO(file_bytes)
                reader = PdfReader(f)
                for idx, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        docs.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "filename": filename,
                                    "page_number": idx + 1,
                                    "document_type": "pdf",
                                    "ingestion_timestamp": timestamp,
                                    "source": filename,
                                },
                            )
                        )
            except Exception as e:
                logger.error("Failed to parse PDF bytes %s: %s", filename, e)
        elif ext == ".docx":
            try:
                f = io.BytesIO(file_bytes)
                doc = DocxDocument(f)
                full_text = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n".join(full_text)
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={
                                "filename": filename,
                                "page_number": 1,
                                "document_type": "docx",
                                "ingestion_timestamp": timestamp,
                                "source": filename,
                            },
                        )
                    )
            except Exception as e:
                logger.error("Failed to parse DOCX bytes %s: %s", filename, e)
        elif ext in (".md", ".markdown", ".txt"):
            try:
                text = file_bytes.decode("utf-8", errors="replace")
                doc_type = "markdown" if ext in (".md", ".markdown") else "txt"
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={
                                "filename": filename,
                                "page_number": 1,
                                "document_type": doc_type,
                                "ingestion_timestamp": timestamp,
                                "source": filename,
                            },
                        )
                    )
            except Exception as e:
                logger.error("Failed to parse text bytes %s: %s", filename, e)
        else:
            logger.warning("Unsupported file type skipped: %s", filename)

        return docs

    # -------------------------------------------------------------- #
    #  Ingestion                                                       #
    # -------------------------------------------------------------- #

    def ingest_directory(self, directory: str) -> int:
        """Scan directory and index PDF, DOCX, TXT, and Markdown files.

        Returns the total number of chunks indexed.
        """
        dir_path = Path(directory)
        if not dir_path.is_dir():
            logger.warning("Data directory not found: %s", directory)
            return 0

        all_docs: List[Document] = []
        supported_extensions = ("*.txt", "*.md", "*.markdown", "*.pdf", "*.docx")
        
        # Scan and parse files
        for ext_pattern in supported_extensions:
            for fpath in dir_path.rglob(ext_pattern):
                parsed_docs = self.parse_file(fpath)
                all_docs.extend(parsed_docs)
                logger.debug("Parsed %s (%d raw doc items)", fpath.name, len(parsed_docs))

        if not all_docs:
            logger.warning("No loadable documents found in %s", directory)
            return 0

        # Chunk documents
        chunks = self.text_splitter.split_documents(all_docs)
        logger.info(
            "Split %d parsed document(s) into %d chunks.", len(all_docs), len(chunks)
        )

        # Ingest to ChromaDB
        self.vector_store.add_documents(chunks)
        
        # Rebuild BM25 search index
        self.rebuild_bm25()
        return len(chunks)

    def ingest_texts(self, texts: List[str], metadatas: Optional[List[dict]] = None) -> int:
        """Ingest raw text strings directly (useful for API text-ingest or testing)."""
        timestamp = datetime.now(timezone.utc).isoformat()
        docs = []
        for i, t in enumerate(texts):
            meta = metadatas[i] if metadatas else {}
            # Standardize required metadata
            meta.setdefault("filename", "raw_text_api")
            meta.setdefault("page_number", 1)
            meta.setdefault("document_type", "txt")
            meta.setdefault("ingestion_timestamp", timestamp)
            meta.setdefault("source", "API")
            docs.append(Document(page_content=t, metadata=meta))

        chunks = self.text_splitter.split_documents(docs)
        self.vector_store.add_documents(chunks)
        
        self.rebuild_bm25()
        logger.info("Ingested %d text(s) → %d chunks.", len(texts), len(chunks))
        return len(chunks)

    def ingest_file_data(self, file_bytes: bytes, filename: str) -> int:
        """Ingest single uploaded file content."""
        parsed_docs = self.parse_file_bytes(file_bytes, filename)
        if not parsed_docs:
            return 0
        chunks = self.text_splitter.split_documents(parsed_docs)
        self.vector_store.add_documents(chunks)
        self.rebuild_bm25()
        return len(chunks)

    # -------------------------------------------------------------- #
    #  Retrieval & Fusion                                              #
    # -------------------------------------------------------------- #

    def hybrid_retrieve(
        self,
        query: str,
        top_k: int = 4,
        filter_dict: Optional[Dict[str, Any]] = None,
    ) -> List[Document]:
        """Perform Hybrid Retrieval combining dense and sparse indices via RRF."""
        k = top_k or self.top_k
        
        # 1. Dense retrieval (ChromaDB)
        # We retrieve double the results to have a rich intersection set
        dense_results = self.vector_store.similarity_search(query, k=k * 2, filter_dict=filter_dict)

        # 2. Sparse retrieval (BM25 keyword search)
        sparse_results: List[Document] = []
        if self.bm25:
            # Simple tokenization
            tokenized_query = query.lower().split()
            scores = self.bm25.get_scores(tokenized_query)
            top_indices = np.argsort(scores)[::-1][:k * 2]
            for idx in top_indices:
                if scores[idx] > 0.0:  # Only matching terms
                    sparse_results.append(self.bm25_docs[idx])

        # 3. Reciprocal Rank Fusion (RRF)
        k_rrf = 60
        rrf_scores: Dict[str, float] = {}
        doc_map: Dict[str, Document] = {}

        def get_doc_key(doc: Document) -> str:
            # Create a unique key for deduplication during rank fusion
            filename = doc.metadata.get("filename", "unknown")
            page = doc.metadata.get("page_number", 1)
            return f"{filename}::pg{page}::{doc.page_content}"

        for rank, doc in enumerate(dense_results):
            key = get_doc_key(doc)
            doc_map[key] = doc
            rrf_scores[key] = rrf_scores.get(key, 0.0) + (1.0 / (k_rrf + rank + 1))

        for rank, doc in enumerate(sparse_results):
            key = get_doc_key(doc)
            doc_map[key] = doc
            rrf_scores[key] = rrf_scores.get(key, 0.0) + (1.0 / (k_rrf + rank + 1))

        # Sort combined results by RRF score descending
        sorted_keys = sorted(rrf_scores.keys(), key=lambda x: rrf_scores[x], reverse=True)
        fused_docs = [doc_map[key] for key in sorted_keys[:k]]

        logger.info(
            "Hybrid retrieve: dense matches=%d, sparse matches=%d, fused total=%d",
            len(dense_results),
            len(sparse_results),
            len(fused_docs),
        )
        return fused_docs

    # -------------------------------------------------------------- #
    #  Query Rewriting                                                 #
    # -------------------------------------------------------------- #

    def rewrite_query(self, query: str, lora_pipeline: Optional[Any] = None) -> str:
        """Rewrite the query to optimize it for keyword & semantic search."""
        if not lora_pipeline or not lora_pipeline.is_loaded:
            logger.debug("Generator pipeline not loaded. Skipping query rewrite.")
            return query

        prompt = (
            "[INST] Task: Rewrite the following user query to make it highly optimized for search engine retrieval. "
            "Rephrase it as a search query containing core keywords and concepts. "
            "Output ONLY the final rephrased search query. Do not add explanations or notes.\n\n"
            f"Query: {query} [/INST]\n"
            "Search Query:"
        )
        try:
            rewritten = lora_pipeline.generate(
                prompt,
                max_new_tokens=40,
                temperature=0.1,
            )
            # Clean output in case model outputs extra quotes or prefixes
            clean_query = rewritten.strip().strip('"').strip("'").replace("\n", " ")
            if clean_query:
                logger.info("Query rewritten: '%s' -> '%s'", query, clean_query)
                return clean_query
        except Exception as e:
            logger.error("Query rewrite failed: %s", e)
            
        return query

    # -------------------------------------------------------------- #
    #  Generation & Citation Formatting                                #
    # -------------------------------------------------------------- #

    def retrieve_and_generate(
        self,
        query: str,
        lora_pipeline: Optional[Any] = None,
        top_k: Optional[int] = None,
        filter_dict: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Perform end-to-end RAG retrieve and synthesis with source citations."""
        k = top_k or self.top_k
        
        # 1. Rewrite query if possible
        search_query = self.rewrite_query(query, lora_pipeline)

        # 2. Retrieve relevant chunks
        retrieved_docs = self.hybrid_retrieve(search_query, top_k=k, filter_dict=filter_dict)

        if not retrieved_docs:
            return {
                "answer": "No relevant context documents were found to answer your question.",
                "sources": [],
            }

        # 3. Generate answer if pipeline is loaded
        if lora_pipeline and lora_pipeline.is_loaded:
            # Build context block
            context_blocks = []
            for i, doc in enumerate(retrieved_docs, 1):
                fname = doc.metadata.get("filename", "unknown")
                page = doc.metadata.get("page_number", 1)
                context_blocks.append(f"[Document: {fname} | Page: {page}]\n{doc.page_content}")
            
            context_str = "\n\n---\n\n".join(context_blocks)

            prompt = (
                "[INST] You are a professional assistant. Answer the user query using ONLY the provided context. "
                "If the context does not contain the answer, state that you cannot answer based on the documents. "
                "Keep the answer factual and clear.\n\n"
                f"Context:\n{context_str}\n\n"
                f"Query: {query} [/INST]\n"
                "Answer:"
            )
            try:
                answer = lora_pipeline.generate(prompt, max_new_tokens=300, temperature=0.3)
            except Exception as e:
                logger.error("Synthesis failed: %s", e)
                answer = f"[Generation Error] Could not synthesize answer. Raw context blocks retrieved:\n\n{context_str}"
        else:
            # Fallback when model is not loaded (e.g. CLI tests / setup)
            context_blocks = []
            for i, doc in enumerate(retrieved_docs, 1):
                fname = doc.metadata.get("filename", "unknown")
                page = doc.metadata.get("page_number", 1)
                context_blocks.append(f"[{fname} - Page {page}]: {doc.page_content}")
            answer = "RAG Pipeline (No Synthesis Model Loaded):\n\n" + "\n\n".join(context_blocks)

        # 4. Extract citations
        sources = []
        seen_sources = set()
        for doc in retrieved_docs:
            filename = doc.metadata.get("filename", "unknown")
            try:
                page = int(doc.metadata.get("page_number", 1))
            except (ValueError, TypeError):
                page = 1
            
            source_key = (filename, page)
            if source_key not in seen_sources:
                seen_sources.add(source_key)
                sources.append({"file": filename, "page": page})

        return {
            "answer": answer.strip(),
            "sources": sources,
        }
